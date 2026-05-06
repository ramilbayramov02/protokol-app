# report.py — Word hesabat generatoru
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
from datetime import datetime
import io

NAVY  = RGBColor(27,  42,  74)
GOLD  = RGBColor(201, 168, 76)
GREEN = RGBColor(22,  163, 74)
RED   = RGBColor(220, 38,  38)
AMBER = RGBColor(217, 119, 6)
GRAY  = RGBColor(100, 116, 139)

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_table(doc, headers: list, rows: list, col_widths: list = None):
    if not rows:
        doc.add_paragraph("Məlumat yoxdur.")
        return
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    # Header
    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        c = hrow.cells[i]
        c.text = h
        run = c.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_bg(c, "1B2A4A")
    # Data rows
    for ri, row_data in enumerate(rows):
        robj = tbl.rows[ri + 1]
        bg   = "DBEAFE" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            c      = robj.cells[ci]
            c.text = str(val) if val is not None else "—"
            run    = c.paragraphs[0].runs[0]
            run.font.size = Pt(9)
            set_cell_bg(c, bg)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row_obj in tbl.rows:
                row_obj.cells[i].width = Inches(w)
    doc.add_paragraph()

def generate_word_report(log_df: pd.DataFrame, delegations: pd.DataFrame,
                          operator: str = "Protokol Xidməti",
                          event_name: str = "VIP Kortej") -> bytes:
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(0.8)
        sec.bottom_margin = Inches(0.8)
        sec.left_margin   = Inches(1.0)
        sec.right_margin  = Inches(1.0)

    def h(text, level=1, color=NAVY):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            run.font.color.rgb = color
        return p

    def p(text, bold=False, color=None, size=10):
        par = doc.add_paragraph()
        run = par.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        return par

    # Statistika
    tot  = len(log_df)
    nok  = (log_df["status"] == "OK").sum()
    ndl  = (log_df["status"] == "Delay").sum()
    npd  = (log_df["status"] == "Pending").sum()
    pok  = round(nok / tot * 100, 1) if tot > 0 else 0
    pdl  = round(ndl / tot * 100, 1) if tot > 0 else 0
    cls  = "GOOD" if pdl == 0 else ("WARNING" if pdl < 20 else "CRITICAL")
    cls_color = GREEN if cls == "GOOD" else (AMBER if cls == "WARNING" else RED)

    # Başlıq
    h("REPUBLIC OF AZERBAIJAN", 1, NAVY)
    h("OFFICE OF THE PROTOCOL SERVICE OF THE PRESIDENT", 2, GRAY)
    doc.add_paragraph()
    h(f"{event_name} — OFFICIAL OPERATIONS REPORT", 1, GOLD)
    p(f"Report Date: {datetime.now().strftime('%d %B %Y')}")
    p(f"Generated at: {datetime.now().strftime('%H:%M:%S')}")
    p(f"Prepared by: {operator}")
    doc.add_paragraph()

    # 1. Xülasə
    h("1. EXECUTIVE SUMMARY", 2, NAVY)
    p(f"Total {len(delegations)} delegations, {tot} events tracked.")
    doc.add_paragraph()
    h(f"Classification: {cls}", 3, cls_color)
    p(f"On time: {nok} ({pok}%)   |   Delayed: {ndl} ({pdl}%)   |   Pending: {npd}", bold=True)
    doc.add_paragraph()

    # 2. Delegasiyalar
    h("2. DELEGATION OVERVIEW", 2, NAVY)
    p(f"{len(delegations)} delegations coordinated through PCC1, PCC2, PCC3.")
    doc.add_paragraph()
    del_rows = []
    for _, d in delegations.sort_values("convoy_order").iterrows():
        del_rows.append([
            d.get("convoy_order",""),
            d.get("country_name",""),
            d.get("leader_name",""),
            d.get("position",""),
            d.get("pcc",""),
            d.get("greeting_location",""),
        ])
    add_table(doc,
        ["#","Country","Leader","Position","PCC","Hotel/Location"],
        del_rows,
        [0.35, 1.2, 1.5, 1.4, 0.6, 1.2])

    # 3. Handshake
    h("3. HANDSHAKE PROTOCOL", 2, NAVY)
    hs_df = log_df[log_df["is_handshake"] == True].copy() if "is_handshake" in log_df.columns else pd.DataFrame()
    if hs_df.empty:
        hs_df = log_df[log_df["event_name"] == "Handshake"].copy()
    hs_rows = []
    for _, r in hs_df.iterrows():
        act = r.get("actual_time","")
        dm  = ""
        if act:
            try:
                pp = [int(x) for x in r["planned_time"].split(":")]
                ap = [int(x) for x in act.split(":")]
                diff = (ap[0]*60+ap[1]) - (pp[0]*60+pp[1])
                dm = f"+{diff}" if diff >= 0 else str(diff)
            except:
                dm = ""
        hs_rows.append([r.get("country_name",""), r.get("pcc",""),
                         r.get("planned_time",""), act or "—", dm or "—", r.get("status","")])
    add_table(doc,
        ["Country","PCC","Planned","Actual","Delta(min)","Status"],
        hs_rows, [1.4, 0.6, 0.9, 0.9, 0.9, 0.9])

    # 4. Gecikmələr
    h("4. DELAY ANALYSIS", 2, NAVY)
    dl_df = log_df[log_df["status"] == "Delay"].copy()
    if dl_df.empty:
        p("No delays recorded.", color=GREEN)
    else:
        p(f"{ndl} events delayed ({pdl}%).")
        doc.add_paragraph()
        dl_rows = []
        for _, r in dl_df.iterrows():
            try:
                pp = [int(x) for x in r["planned_time"].split(":")]
                ap = [int(x) for x in r["actual_time"].split(":")]
                diff = (ap[0]*60+ap[1]) - (pp[0]*60+pp[1])
                dm = str(diff)
            except:
                dm = "—"
            dl_rows.append([r.get("country_name",""), r.get("event_name",""),
                             r.get("planned_time",""), r.get("actual_time",""),
                             dm, r.get("delay_reason","—"), r.get("recorded_by","—")])
        add_table(doc,
            ["Country","Event","Planned","Actual","Delay(min)","Reason","PCC"],
            dl_rows, [1.1, 1.3, 0.8, 0.8, 0.8, 1.1, 0.6])

    # 5. PCC
    h("5. PCC PERFORMANCE", 2, NAVY)
    pcc_rows = []
    for pcc_code in ["PCC1","PCC2","PCC3"]:
        sub  = log_df[log_df["pcc"] == pcc_code]
        dels = sub["country_name"].nunique()
        tot2 = len(sub)
        ok2  = (sub["status"] == "OK").sum()
        dl2  = (sub["status"] == "Delay").sum()
        rate = f"{round(ok2/tot2*100,1)}%" if tot2 > 0 else "—"
        pcc_rows.append([pcc_code, dels, tot2, ok2, dl2, rate])
    add_table(doc,
        ["PCC","Delegations","Total Events","On Time","Delayed","Rate"],
        pcc_rows, [0.7, 0.9, 1.1, 0.9, 0.9, 0.8])

    # 6. Yekun
    h("6. FINAL ASSESSMENT", 2, NAVY)
    h(f"Classification: {cls}", 3, cls_color)
    cls_txt = {
        "GOOD":     "All protocol events executed within acceptable parameters.",
        "WARNING":  "Minor delays recorded. Protocol integrity maintained.",
        "CRITICAL": "Significant delays impacted execution. Formal review required.",
    }[cls]
    p(cls_txt)
    doc.add_paragraph()
    p("This report was generated automatically by the VIP Convoy Control Tower System "
      "of the Protocol Service of the President of the Republic of Azerbaijan.",
      color=GRAY, size=9)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
